import os
from datetime import datetime

import openpyxl
import subprocess
from docx import Document
from docx2pdf import convert
from concurrent.futures import ThreadPoolExecutor, as_completed

from app.state.app_state import app_state


# Function to replace placeholder in paragraphs
def replace_placeholder_in_paragraphs(doc, placeholder, replacement):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)


# Function to replace placeholder in tables
def replace_placeholder_in_tables(doc, placeholder, replacement):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, replacement)


# Function to replace placeholder in headers and footers
def replace_placeholder_in_headers_footers(doc, placeholder, replacement):
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)

        footer = section.footer
        for paragraph in footer.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)


# Function to replace placeholder in footnotes and endnotes (does not currently work)
def replace_placeholder_in_notes(doc, placeholder, replacement):
    for footnote in doc.footnotes:
        if placeholder in footnote.text:
            footnote.text = footnote.text.replace(placeholder, replacement)
    for endnote in doc.endnotes:
        if placeholder in endnote.text:
            endnote.text = endnote.text.replace(placeholder, replacement)


# Function to replace placeholder in entire document
def replace_placeholder_in_document(doc, placeholder, replacement):
    replace_placeholder_in_paragraphs(doc, placeholder, replacement)
    replace_placeholder_in_tables(doc, placeholder, replacement)
    replace_placeholder_in_headers_footers(doc, placeholder, replacement)


def get_data_as_list(data):
    data_list = []
    for row in data.iter_rows(values_only=True):
        data_row = []
        for cell in row:
            if (cell is not None) and (cell != ""):
                if isinstance(cell, str):
                    data_row.append(cell.strip())
                else:
                    data_row.append(str(cell).strip())
            else:
                data_row.append(cell)
        data_list.append(data_row)

    return data_list


def is_empty(data):
    if (data is None) or (data == "None") or (data == ""):
        return True
    else:
        return False


def convert_to_pdf(output_folder_path, output_file_name):
    convert(os.path.join(output_folder_path, output_file_name))
    os.remove(os.path.join(output_folder_path, output_file_name))


def convert_all_docx_to_pdf(output_folder_path):
    if app_state.save_as_pdf.get():
        for file in os.listdir(output_folder_path):
            if file.endswith(".docx"):
                print(f"Converting {file} to PDF...")
                convert_to_pdf(output_folder_path, file)


def process_row(row, template_file_path, doc_name):
    # Open the Word document
    doc = Document(template_file_path)

    date = datetime.today()

    placeholder = "[DATE]"
    replacement = date.strftime("%d %B %Y")

    replace_placeholder_in_document(doc, placeholder, replacement)

    personalized_name_list = [row[i] for i in app_state.get_naming_columns()]
    personalized_name = ""

    for i, word in enumerate(personalized_name_list):
        if is_empty(word):
            word = "N/A"
        
        personalized_name += word
        if i < len(personalized_name_list) - 1:
            personalized_name += " "


    for placeholder_and_replacement_column in app_state.get_placeholders_and_replacement_columns():

        placeholder = placeholder_and_replacement_column[0]
        replacement_column = placeholder_and_replacement_column[1]
        replacement = row[replacement_column]

        if is_empty(replacement):
            replacement = "N/A"
        
        replace_placeholder_in_document(doc, placeholder, replacement)

        # Save the modified document
        os.makedirs(app_state.output_folder_path, exist_ok=True)

        doc.save(f"{app_state.output_folder_path}\\{personalized_name}_{doc_name}_{date.year}.docx")


def toggle_processing_state():
    app_state.is_processing = not app_state.is_processing
    if app_state.is_processing:
        app_state.propogate_template_button.configure(state="disabled")
    elif not app_state.is_processing:
        app_state.propogate_template_button.configure(state="normal")


def propogate_template(excel_file_path, template_file_path):
    toggle_processing_state()

    wb = openpyxl.load_workbook(excel_file_path, data_only=True, read_only=True)
    
    excel_sheet_name = app_state.get_sheet_name()
    excel_sheet = get_data_as_list(wb[excel_sheet_name])
    doc_name = app_state.get_document_name()

    # Implemnent multithreading for processing rows
    # max_workers = max(1, os.cpu_count() - 1)

    # with ThreadPoolExecutor(max_workers=max_workers) as executor:
    #     futures = []
    #     for row in excel_sheet:
    #         futures.append(executor.submit(process_row, row, template_file_path, doc_name))
        
    #     for future in as_completed(futures):
    #         try:
    #             future.result()
    #         except Exception as e:
    #             print(f"Error processing row: {e}")

    for row in excel_sheet:
        try:
            process_row(row, template_file_path, doc_name)
        except Exception as e:
            print(f"Error processing row {row}: {e}")

    convert_all_docx_to_pdf(app_state.output_folder_path)

    toggle_processing_state()